# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


import json
import pandas as pd
import numpy as np
import os
import re
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.dml import MSO_THEME_COLOR
import glob
from tqdm import tqdm

#This function divide case/item stem into non_html and html component
def Extracthtml(string):
    #Detect html
    if r'<table' in string:
        tablestr=string[string.find('<table'):]
        htmltables=pd.read_html(tablestr)
        return string[:string.find('<br')], htmltables
    else:
        return string.replace('<sub>','').replace('</sub>',''), None
#This function set column width for Microsoft Word Table
def setcolwidth(col, value):
    for cell in col.cells:
        cell.width = Cm(value)
    return col

#This function adds a pandas dataframe as a table in microsoft word document
def pd2word(doc, df, header=True, style='Light List Accent 1', cwidth=None):
    if header==True:

        # add a table to the end and create a reference variable
        # extra row is so we can add the header row
        t = doc.add_table(df.shape[0] + 1, df.shape[1])

        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i +1, j).text = str(df.values[i, j])
    else:
        t = doc.add_table(df.shape[0], df.shape[1])

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i , j).text = str(df.values[i, j])
    t.style = style
    if cwidth!=None:
        t.allow_autofit =False
        t.autofit = False
        for p in range(len(cwidth)):
            setcolwidth(t.columns[p],cwidth[p])

    return doc

#This class extract information from exampackage and marking json file (filepath), and convert the extraction into a microsoft word file
class case_report():
    def __init__(self, filepath, exampackage):
        self.filepath=filepath
        with open(self.filepath, 'rb') as json_file:
            self.jsonf = json.load(json_file)
        with open(exampackage, 'rb') as exam_file:
            self.MOC5 = json.load(exam_file)
        self.Nmarkers=len(self.jsonf['Markers'])
        self.Nquestions=len(self.jsonf['Questions'])
        #convert MOC5 into a case dictionary
        self.casedic={}
        for case in self.MOC5['TdmItems']['$values']:
            cdic={}
            cdic['CaseStem']=case['CaseStem']['Stem']
            idic={}
            for question in case['Questions']['$values']:
                idic[question['Id']]=question

            cdic['Questions']=idic
            self.casedic[case['Id']]=cdic
        # Basic Info
        self.CaseId=int(self.jsonf['CaseNumber'])
        basicTable = pd.DataFrame()
        basicTable['c1'] = ['Basic Information','Candidate Code:', 'Exam Session:', 'Case:']
        basicTable['c2'] = ['', self.jsonf['CandidateCode'], '2020-S2-Oct',
                            self.CaseId]

        # Case Stem
        Cstem = self.casedic[self.CaseId]['CaseStem'].replace('&nbsp;', ' ')
        Cstem_text, Cstem_table = Extracthtml(Cstem)
        # Generate word document
        document = Document()
        document.add_heading('Rescore Report', 0)
        document.add_heading('Selection Criteria', 1)
        document = pd2word(document, basicTable, header=False, cwidth=[7.5,7.5])
        document.add_heading('Case Stem', 1)
        document.add_paragraph(Cstem_text)
        if Cstem_table !=None:
            for tb in Cstem_table:
                if tb.columns[0]==0:
                    document = pd2word(document, tb.fillna(''), header=False, style='Light Grid Accent 1')
                else:
                    document = pd2word(document, tb.fillna(''), style='Light Grid Accent 1')

        for i in range(self.Nquestions):
            if i+1 in self.casedic[self.CaseId]['Questions'].keys():
                Questioni = self.jsonf['Questions'][i]
                QuestioniMOC5 = self.casedic[self.CaseId]['Questions'][i+1]
                ###Generate a report for each question

                #Question Stem
                Qstem=QuestioniMOC5['QuestionWriteIn']['LeadingQuestion'].replace('&nbsp;', ' ')
                Qstem_text, Qstem_table = Extracthtml(Qstem)
                #Answer Key Table
                generalScoringRule=pd.DataFrame({
                    'c1':['Rules', 'Notes', 'MaximumAllowableResponses', 'MaximumAllowableScore', 'Rationale'],
                    'c2': ['', QuestioniMOC5['ScoringNotes'], QuestioniMOC5['QuestionWriteIn']['MaximumNumberOfAnswers'], QuestioniMOC5['QuestionWriteIn']['MaximumScore'], Questioni['ScoringKey']['Rationale']]
                })

                Answerkeys=pd.DataFrame(QuestioniMOC5['QuestionWriteIn']['CorrectAnswers']['$values'])[['Answer', 'Synonyms', 'NotAcceptable', 'Score']]

                #Remove html char
                Answerkeys['Answer']=Answerkeys['Answer'].str.replace('&nbsp;', ' ')
                Answerkeys['Synonyms']=Answerkeys['Synonyms'].str.replace('&nbsp;', ' ')
                Answerkeys['NotAcceptable'] = Answerkeys['NotAcceptable'].str.replace('&nbsp;', ' ')
                AnswerkeysCorrect=Answerkeys.loc[Answerkeys['Score']>0]
                AnswerkeysIncorrect = Answerkeys.loc[Answerkeys['Score'] == 0]
                #Response Score Table

                ResponseTable=pd.DataFrame(Questioni['Responses'])
                if ResponseTable.shape!=(0,0):
                    ResponseTable['Response']=ResponseTable['Response'].str.replace('&nbsp;', ' ')

                    ScoreTable=pd.DataFrame(self.jsonf['Markers'][-1]['QuestionMarks'][i]['ResponseMarks'])
                    ResponseScore=pd.DataFrame(ScoreTable[['Mark', 'NumberOfResponses','Score','Notes','ExceededMax']])
                    ResponseScore.columns=['Mark', 'Number of Responses','Score','Notes','Exceeded Max']
                    ResponseScore['Notes']=ResponseScore['Notes'].str.replace('&nbsp;', ' ')
                    ResponseScore.insert(0, 'Candidate Response', ResponseTable['Response'])
                    #self.ResponseScore[['Revised Mark', 'Revised Number of Responses', 'Revised Score', 'Remark Notes', 'Revised ExceededMax']]=None

                    #Revised Score Table
                    Revision=pd.DataFrame()
                    Revision['Candidate Response']=ResponseTable['Response']
                    Revision['Revised Number of Responses']=''
                    Revision['Revised Score']=''
                    Revision['Remark Notes'] = ''
                    Revision['Revised Exceeded Max'] = ''
                else:
                    ResponseScore=pd.DataFrame({'Candidate Response':[''], 'Mark':[''], 'Number of Responses':[''],'Score':[''],'Notes':[''],'Exceeded Max':['']})
                    Revision=pd.DataFrame({'Candidate Response':[''],'Revised Number of Responses':[''],'Revised Score':[''], 'Remark Notes':[''], 'Revised Exceeded Max':['']})
                #Signature and comments
                Signature=pd.DataFrame({'c1':['Signature', 'Comments'], 'c2':['', '']})


                document.add_heading('Question Stem', 1)
                document.add_heading(f'Question {i+1}', 3)
                document.add_paragraph(Qstem_text)
                if Qstem_table != None:
                    for tb in Qstem_table:
                        if tb.columns[0] == 0:
                            document = pd2word(document, tb.fillna(''), header=False, style='Light Grid Accent 1')
                        else:
                            document = pd2word(document, tb.fillna(''), style='Light Grid Accent 1')
                document.add_heading('General Scoring Rules', 1)
                document = pd2word(document, generalScoringRule.fillna(''), header=False, cwidth=[7.5,7.5])
                document.add_heading('Candidate Responses and Scores', 1)
                document = pd2word(document, ResponseScore.fillna(''), cwidth=[4.5,1.5,2.5,1.5,2.5,2.5])
                document.add_heading('Scoring keys', 1)
                document = pd2word(document, AnswerkeysCorrect.fillna(''), cwidth=[4.5,4.5,4.5,1.5])
                if AnswerkeysIncorrect.shape[0]>0:
                    run = document.add_heading().add_run('Unacceptable Answers')
                    font = run.font
                    font.color.theme_color = MSO_THEME_COLOR.ACCENT_2
                    #document.add_heading('Unacceptable Answers', 2)
                    document = pd2word(document, AnswerkeysIncorrect[['Answer', 'Synonyms', 'Score']].fillna(''), style='Light List Accent 2')

                document.add_heading('Score Revision', 1)
                document = pd2word(document, Revision.fillna(''), cwidth=[5, 2.5, 2.5, 2.5, 2.5], style='Light Grid Accent 1')
                document.add_heading('Comments and Signature', 1)
                document = pd2word(document, Signature, header=False, style='Light Grid Accent 1', cwidth=[5, 10])
        document.save(f"C:/Users/qguo/Documents/TDM_recheck/"+filepath[:-5]+".docx")





# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    exampackage = r'\\MCC-DFS01\Mcc.Tdm.Marking\PROD\2021-Jan-Source File Backup files\ExamSnapshot-PRA-2020-TF2.v6.01.json'
    os.chdir(r'\\mcc-dfs01\databases\Merlin\Examinations\TDM\2020\Form2\Raw Data\MarkingJsonFiles\markingfiles')
    MarkingFiles = pd.DataFrame({'FileName': glob.glob("*.json")})
    MarkingFiles['CandidateCode'] = MarkingFiles['FileName'].str.extract(r'-(\d+)')[0].astype(int).tolist()
    requestIDs = [1131963751, 1531888350]
    RequestFiles = MarkingFiles.loc[MarkingFiles['CandidateCode'].isin(requestIDs), 'FileName'].tolist()

    for file in tqdm(RequestFiles):
        gg = case_report(file, exampackage)

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
